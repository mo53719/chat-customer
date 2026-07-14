"""知识库管理路由。

解析层职责：把上传的文件解析成 List[ChunkDraft]（结构化切片 + 元数据），
返回给 knowledge_service 入库。文本类返回带分隔的整文 + source 元数据，
由 Ingestor 二次切分；表格类直接逐行产出独立 chunk，每条带 sheet/row。
"""
from __future__ import annotations

import asyncio
import codecs
import csv
import hashlib
import io
import json
from dataclasses import dataclass, field
from typing import Any

from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, WebSocket, WebSocketDisconnect, Query

from app.services.knowledge_service import knowledge_service
from app.storage.sqlite.repositories.knowledge_repo import knowledge_repo
from app.storage.qdrant.ingestor import Ingestor
from ..schemas.common import ApiResponse
from ..deps import get_current_user

router = APIRouter(prefix="/api/knowledge", tags=["知识库"])

# 单文件最大 50MB，防止 raw = await file.read() 把内存撑爆
MAX_UPLOAD_BYTES = 50 * 1024 * 1024

# 允许上传的文件扩展名
ALLOWED_EXTS = ("xlsx", "xlsm", "xls", "csv", "txt", "md", "json", "log", "pdf", "docx")

# ── 进度管理器（内存） ────────────────────────────────────────
_progress: dict[str, dict[str, Any]] = {}
_ws_clients: dict[str, list[WebSocket]] = {}


async def _broadcast_progress(doc_id: str, phase: str, current: int, total: int, message: str):
    """更新进度并推送给所有在听的 WebSocket 客户端。"""
    pct = round(current / total * 100, 1) if total > 0 else 0
    payload = {
        "doc_id": doc_id,
        "phase": phase,
        "current": current,
        "total": total,
        "message": message,
        "pct": pct,
    }
    _progress[doc_id] = payload
    dead: list[WebSocket] = []
    for ws in _ws_clients.get(doc_id, []):
        try:
            await ws.send_json(payload)
        except Exception:
            dead.append(ws)
    for ws in dead:
        _ws_clients.get(doc_id, []).remove(ws)


async def _process_doc(doc_id: str, title: str, drafts: list, file_type: str,
                       source: str, uploaded_by: str, chunk_size: int, overlap: int,
                       file_hash: str | None = None, parent_doc_id: str | None = None,
                       is_current: bool = True, total_pages: int | None = None,
                       uploader_role: str | None = None, char_count: int = 0):
    """后台处理：向量化 + 入库，同时推送进度。复用路由层已创建的 doc_id。"""
    async def _cb(phase, cur, tot, msg):
        await _broadcast_progress(doc_id, phase, cur, tot, msg)

    await knowledge_service.upload_drafts(
        doc_id=doc_id, title=title, drafts=drafts, file_type=file_type,
        source=source, uploaded_by=uploaded_by,
        chunk_size=chunk_size, overlap=overlap,
        progress_callback=_cb,
        file_hash=file_hash, parent_doc_id=parent_doc_id,
        is_current=is_current, total_pages=total_pages,
        uploader_role=uploader_role, char_count=char_count,
    )


@dataclass
class ChunkDraft:
    """解析器产出的最小切片单元。"""
    content: str
    source: str = ""
    sheet: str | None = None
    row: int | None = None
    heading_path: str | None = None   # 例如 "Sheet1 > 退费政策 > 时效"
    extra: dict[str, Any] = field(default_factory=dict)
    page_no: int | None = None        # PDF/Word 页码
    section: str | None = None        # Word 章节标题
    char_start: int | None = None     # 原文起始偏移
    char_end: int | None = None       # 原文结束偏移
    char_count: int | None = None     # 字符数


def _decode_text(raw: bytes) -> str:
    """智能识别文本编码（utf-8 / utf-8-sig / gbk / gb18030）。"""
    if not raw:
        return ""
    if raw[:3] == b"\xef\xbb\xbf":
        return raw[3:].decode("utf-8", errors="ignore")
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    for enc in ("gbk", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _parse_xlsx(raw: bytes, source: str) -> list[ChunkDraft]:
    """从 .xlsx / .xlsm 按行产出 ChunkDraft。处理合并单元格表头。"""
    from openpyxl import load_workbook
    from openpyxl.utils import range_boundaries
    wb = load_workbook(io.BytesIO(raw), read_only=False, data_only=True)
    drafts: list[ChunkDraft] = []

    for sheet in wb.worksheets:
        merged_values: dict[tuple[int, int], Any] = {}
        if hasattr(sheet, "merged_cells") and sheet.merged_cells.ranges:
            for mr in sheet.merged_cells.ranges:
                min_col, min_row, max_col, max_row = range_boundaries(str(mr))
                top_val = sheet.cell(row=min_row, column=min_col).value
                for r in range(min_row, max_row + 1):
                    for c in range(min_col, max_col + 1):
                        merged_values[(r, c)] = top_val

        rows_iter = sheet.iter_rows(values_only=False)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            continue
        headers: list[str] = []
        for cell in header_row:
            v = merged_values.get((cell.row, cell.column), cell.value)
            headers.append(str(v).strip() if v is not None else f"col{cell.column}")

        for row in rows_iter:
            if all((c.value is None or (isinstance(c.value, str) and not c.value.strip()))
                   and (c.row, c.column) not in merged_values for c in row):
                continue
            cells = []
            for cell in row:
                v = merged_values.get((cell.row, cell.column), cell.value)
                if v is None:
                    continue
                s = str(v).strip()
                if not s:
                    continue
                idx = cell.column - 1
                h = headers[idx] if idx < len(headers) else f"col{cell.column}"
                cells.append(f"{h}: {s}")
            if cells:
                drafts.append(ChunkDraft(
                    content=" | ".join(cells),
                    source=source,
                    sheet=sheet.title,
                    row=row[0].row,
                    heading_path=sheet.title,
                ))
    wb.close()
    return drafts


def _parse_xls(raw: bytes, source: str) -> list[ChunkDraft]:
    """从老格式 .xls 按行解析（要求 xlrd<2.0.0）。"""
    try:
        import xlrd
    except ImportError:
        raise HTTPException(status_code=400, detail=".xls 暂不支持，请另存为 .xlsx 后再上传")
    if hasattr(xlrd, "__version__") and not xlrd.__version__.startswith("1."):
        raise HTTPException(status_code=400, detail=".xls 需要 xlrd<2.0.0，请联系管理员")
    book = xlrd.open_workbook(file_contents=raw)
    drafts: list[ChunkDraft] = []
    for sheet in book.sheets():
        if sheet.nrows == 0:
            continue
        headers = [str(sheet.cell_value(0, c)).strip() or f"col{c}" for c in range(sheet.ncols)]
        for r in range(1, sheet.nrows):
            row = sheet.row_values(r)
            if all(c == "" or c is None for c in row):
                continue
            cells = []
            for h, v in zip(headers, row):
                if v == "" or v is None:
                    continue
                cells.append(f"{h}: {v}")
            if cells:
                drafts.append(ChunkDraft(
                    content=" | ".join(cells),
                    source=source,
                    sheet=sheet.name,
                    row=r + 1,
                    heading_path=sheet.name,
                ))
    return drafts


def _parse_csv(raw: bytes, source: str) -> list[ChunkDraft]:
    """CSV 按行解析，每行一个 chunk。"""
    text = _decode_text(raw)
    reader = csv.reader(io.StringIO(text))
    try:
        headers = next(reader)
    except StopIteration:
        return []
    headers = [h.strip() or f"col{i}" for i, h in enumerate(headers)]
    drafts: list[ChunkDraft] = []
    for r_idx, row in enumerate(reader, start=2):
        if not row or all(not c.strip() for c in row):
            continue
        cells = []
        for h, v in zip(headers, row):
            if not v or not v.strip():
                continue
            cells.append(f"{h}: {v.strip()}")
        if cells:
            drafts.append(ChunkDraft(
                content=" | ".join(cells),
                source=source,
                sheet=None,
                row=r_idx,
                heading_path=source,
            ))
    return drafts


def _parse_plain(raw: bytes, source: str) -> list[ChunkDraft]:
    """txt/md/json/log 整文返回，由 Ingestor 二次切分。"""
    text = _decode_text(raw)
    if not text.strip():
        return []
    return [ChunkDraft(content=text, source=source, heading_path=source,
                       char_count=len(text))]


def _parse_pdf(raw: bytes, source: str) -> list[ChunkDraft]:
    """PDF 解析：按页提取文字，每页产出带 page_no 的 ChunkDraft。

    去页眉页脚：连续两页完全一样的首行/尾行直接丢掉。
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(status_code=400, detail="PDF 解析需要 pypdf 库，请联系管理员安装")

    reader = PdfReader(io.BytesIO(raw))
    drafts: list[ChunkDraft] = []
    total_pages = len(reader.pages)

    if total_pages == 0:
        return []

    # 收集所有页的文本和首尾行
    all_texts: list[str] = []
    all_first_lines: list[str] = []
    all_last_lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if not text.strip():
            all_texts.append("")
            all_first_lines.append("")
            all_last_lines.append("")
            continue
        lines = text.split("\n")
        all_texts.append(text)
        all_first_lines.append(lines[0].strip() if lines else "")
        all_last_lines.append(lines[-1].strip() if lines else "")

    # 检测重复的页眉页脚（连续两页相同即视为页眉/页脚）
    for i, text in enumerate(all_texts):
        if not text.strip():
            continue
        lines = text.split("\n")
        # 去掉页眉（首行与前后页相同）
        if i > 0 and all_first_lines[i] and all_first_lines[i] == all_first_lines[i - 1]:
            if lines and lines[0].strip() == all_first_lines[i]:
                lines.pop(0)
        elif i < total_pages - 1 and all_first_lines[i] and all_first_lines[i] == all_first_lines[i + 1]:
            if lines and lines[0].strip() == all_first_lines[i]:
                lines.pop(0)
        # 去掉页脚（尾行与前后页相同）
        if i > 0 and all_last_lines[i] and all_last_lines[i] == all_last_lines[i - 1]:
            if lines and lines[-1].strip() == all_last_lines[i]:
                lines.pop()
        elif i < total_pages - 1 and all_last_lines[i] and all_last_lines[i] == all_last_lines[i + 1]:
            if lines and lines[-1].strip() == all_last_lines[i]:
                lines.pop()

        clean = "\n".join(lines).strip()
        if clean:
            drafts.append(ChunkDraft(
                content=clean,
                source=source,
                page_no=i + 1,
                heading_path=f"{source} 第{i + 1}页",
                char_count=len(clean),
            ))

    if not drafts:
        raise HTTPException(
            status_code=400,
            detail="该 PDF 无文本层（可能是扫描件），请先用 OCR 工具转 Word 后再上传",
        )
    return drafts


def _parse_docx(raw: bytes, source: str) -> list[ChunkDraft]:
    """Word 解析：按段落提取，识别标题样式（Heading 1/2/3）构建章节路径。

    每段产出带 heading_path 的 ChunkDraft，由 Ingestor 二次切分。
    """
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=400, detail="Word 解析需要 python-docx 库，请联系管理员安装")

    doc = Document(io.BytesIO(raw))
    drafts: list[ChunkDraft] = []
    # 当前章节路径栈
    heading_stack: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else ""
        # 判断是否为标题
        if style.startswith("Heading"):
            level = 1
            try:
                level = int(style.replace("Heading", "").strip())
            except ValueError:
                pass
            # 更新章节路径栈
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(text)
            continue

        # 普通段落
        heading_path = " > ".join(heading_stack) if heading_stack else source
        drafts.append(ChunkDraft(
            content=text,
            source=source,
            heading_path=heading_path,
            section=heading_stack[-1] if heading_stack else None,
            char_count=len(text),
        ))

    if not drafts:
        # 如果没有段落，尝试提取所有文本
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if full_text:
            drafts.append(ChunkDraft(
                content=full_text,
                source=source,
                heading_path=source,
                char_count=len(full_text),
            ))

    if not drafts:
        raise HTTPException(status_code=400, detail="Word 文件无文本内容")

    return drafts


def _extract_chunks(raw: bytes, ext: str, source: str) -> list[ChunkDraft]:
    """根据扩展名分发到不同解析器，统一返回 ChunkDraft 列表。"""
    ext = (ext or "").lower().lstrip(".")
    if ext in ("xlsx", "xlsm"):
        return _parse_xlsx(raw, source)
    if ext == "xls":
        return _parse_xls(raw, source)
    if ext == "csv":
        return _parse_csv(raw, source)
    if ext == "pdf":
        return _parse_pdf(raw, source)
    if ext == "docx":
        return _parse_docx(raw, source)
    # txt / md / json / log
    return _parse_plain(raw, source)


def _ext_of(filename: str | None) -> str:
    if not filename or "." not in filename:
        return "txt"
    return filename.rsplit(".", 1)[-1].lower()


# ── API 端点 ──────────────────────────────────────────────────

@router.get("", response_model=ApiResponse)
async def list_docs(user: dict = Depends(get_current_user)):
    return ApiResponse(data=await knowledge_service.list_docs())


@router.websocket("/ws/{doc_id}")
async def ws_progress(websocket: WebSocket, doc_id: str):
    """WebSocket 进度推送：前端连接后接收实时向量化进度。"""
    await websocket.accept()
    _ws_clients.setdefault(doc_id, []).append(websocket)
    if doc_id in _progress:
        await websocket.send_json(_progress[doc_id])
    try:
        while True:
            await websocket.receive_text()
    except (WebSocketDisconnect, Exception):
        pass
    finally:
        _ws_clients.get(doc_id, []).remove(websocket)


@router.get("/config", response_model=ApiResponse)
async def get_config(user: dict = Depends(get_current_user)):
    return ApiResponse(data=await knowledge_repo.get_config())


@router.post("/config", response_model=ApiResponse)
async def save_config(data: dict[str, Any], user: dict = Depends(get_current_user)):
    await knowledge_repo.save_config(
        chunk_size=int(data.get("chunk_size", 400)),
        overlap=int(data.get("overlap", 80)),
    )
    return ApiResponse(message="配置已保存")


@router.post("/upload", response_model=ApiResponse)
async def upload(file: UploadFile = File(...),
                 title: str = Form(None),
                 chunk_size: int = Form(400),
                 overlap: int = Form(80),
                 action: str = Form("new"),  # new / rebuild / new_version / cancel
                 parent_doc_id: str = Form(None),
                 user: dict = Depends(get_current_user)):
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="文件为空")
    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="文件过大，最大支持 50MB")

    ext = _ext_of(file.filename)
    if ext not in ALLOWED_EXTS:
        raise HTTPException(
            status_code=400,
            detail=f"暂不支持 .{ext} 格式（仅支持 {' / '.join(ALLOWED_EXTS)}）",
        )

    # 计算文件指纹
    file_hash = hashlib.sha256(raw).hexdigest()

    try:
        drafts = _extract_chunks(raw, ext, source=file.filename or "")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件解析失败：{e}")

    if not drafts:
        raise HTTPException(status_code=400, detail="文件解析后无可用内容")

    title = title or file.filename or "未命名"
    import uuid as _uuid

    # 计算总页数和总字数
    total_pages = None
    if ext == "pdf":
        total_pages = len(drafts)
    char_count = sum(getattr(d, "char_count", 0) or 0 for d in drafts) or len(raw)

    # ── 去重检查 ──
    conv = ""
    existing = await knowledge_repo.find_by_hash(file_hash)
    if existing:
        if action == "rebuild":
            # 重建索引：删除旧向量，元数据复用
            await Ingestor.delete_doc(existing.doc_id)
            doc_id = existing.doc_id
            await knowledge_repo.update_status(doc_id, "processing")
            await knowledge_repo.update_meta(doc_id, char_count=char_count, total_pages=total_pages)
            asyncio.create_task(_process_doc(
                doc_id=doc_id, title=title, drafts=drafts, file_type=ext,
                source=file.filename or "", uploaded_by=user["username"],
                chunk_size=chunk_size, overlap=overlap,
                file_hash=file_hash, parent_doc_id=existing.parent_doc_id,
                is_current=True, total_pages=total_pages,
                uploader_role=user.get("role"), char_count=char_count,
            ))
            return ApiResponse(data={"doc_id": doc_id, "status": "processing", "action": "rebuild"})
        elif action == "new_version":
            # 存为新版本：旧版失效
            await knowledge_repo.mark_not_current(existing.doc_id)
            version_no = await knowledge_repo.get_next_version_no(existing.parent_doc_id or existing.doc_id)
            conv = f"v{version_no}"
            parent_doc_id = existing.parent_doc_id or existing.doc_id
        elif action == "cancel":
            return ApiResponse(data={"doc_id": None, "status": "cancelled", "message": "已取消上传"})
        else:
            # action == "new" 但已有重复 → 返回提示让前端弹窗
            return ApiResponse(data={
                "doc_id": None,
                "status": "duplicate",
                "existing_doc": {
                    "doc_id": existing.doc_id,
                    "title": existing.title,
                    "source": existing.source,
                    "created_at": existing.created_at,
                },
                "message": f"文档「{existing.title}」已存在（文件指纹相同），请选择处理方式",
            })

    # 新建文档
    doc_id = f"doc_{_uuid.uuid4().hex[:16]}"
    await knowledge_repo.create(
        doc_id=doc_id, title=title, source=file.filename, file_type=ext,
        uploaded_by=user["username"], file_hash=file_hash,
        parent_doc_id=parent_doc_id, is_current=1,
        embedding_model=None,  # 由 service 层填充
        total_pages=total_pages, uploader_role=user.get("role"),
        char_count=char_count,
    )

    # 写版本记录
    if parent_doc_id:
        version_no = await knowledge_repo.get_next_version_no(parent_doc_id)
        await knowledge_repo.create_version_row(
            doc_id=doc_id, parent_doc_id=parent_doc_id,
            version_no=version_no, file_hash=file_hash,
            chunk_count=0, note=conv,
        )
    else:
        await knowledge_repo.create_version_row(
            doc_id=doc_id, parent_doc_id=None,
            version_no=1, file_hash=file_hash,
            chunk_count=0, note="初始版本",
        )

    asyncio.create_task(_process_doc(
        doc_id=doc_id, title=title, drafts=drafts, file_type=ext,
        source=file.filename or "", uploaded_by=user["username"],
        chunk_size=chunk_size, overlap=overlap,
        file_hash=file_hash, parent_doc_id=parent_doc_id,
        is_current=True, total_pages=total_pages,
        uploader_role=user.get("role"), char_count=char_count,
    ))
    return ApiResponse(data={"doc_id": doc_id, "status": "processing"})


@router.delete("/{doc_id}", response_model=ApiResponse)
async def delete_doc(doc_id: str, user: dict = Depends(get_current_user)):
    await knowledge_service.delete(doc_id, deleted_by=user["username"])
    return ApiResponse(message="已删除")


# ── 溯源查询端点 ──────────────────────────────────────────────

@router.get("/{doc_id}/chunks", response_model=ApiResponse)
async def get_chunks(doc_id: str,
                     limit: int = Query(100, ge=1, le=500),
                     user: dict = Depends(get_current_user)):
    """查看某文档的所有知识片段（从 Qdrant 滚动获取）。"""
    try:
        chunks = await Ingestor.scroll_chunks(doc_id, limit=limit)
        return ApiResponse(data={"doc_id": doc_id, "chunks": chunks, "total": len(chunks)})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取切片失败：{e}")


@router.get("/by-hash/{file_hash}", response_model=ApiResponse)
async def get_by_hash(file_hash: str, user: dict = Depends(get_current_user)):
    """按文件指纹查找已有文档。"""
    meta = await knowledge_repo.find_by_hash(file_hash)
    if not meta:
        return ApiResponse(data=None, message="未找到该文件指纹对应的文档")
    return ApiResponse(data=meta.to_dict())


@router.get("/{doc_id}/versions", response_model=ApiResponse)
async def get_versions(doc_id: str, user: dict = Depends(get_current_user)):
    """查看某文档的所有历史版本。"""
    meta = await knowledge_repo.get(doc_id)
    if not meta:
        raise HTTPException(status_code=404, detail="文档不存在")
    parent = meta.parent_doc_id or meta.doc_id
    versions = await knowledge_repo.list_versions(parent)
    return ApiResponse(data={"doc_id": doc_id, "versions": versions})


@router.post("/{doc_id}/restore-version", response_model=ApiResponse)
async def restore_version(doc_id: str,
                          version_no: int = Query(..., ge=1),
                          user: dict = Depends(get_current_user)):
    """将指定版本恢复为当前生效版本。"""
    meta = await knowledge_repo.get(doc_id)
    if not meta:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 找到版本链的根
    parent = meta.parent_doc_id or meta.doc_id
    # 找到目标版本的 doc_id
    rows = await knowledge_repo.list_versions(parent)
    target_doc_id = None
    for r in rows:
        if r["version_no"] == version_no:
            target_doc_id = r["doc_id"]
            break

    if not target_doc_id:
        raise HTTPException(status_code=404, detail=f"版本 v{version_no} 不存在")

    # 当前所有生效版本失效
    all_versions = await knowledge_repo.list_versions(parent)
    for v in all_versions:
        await knowledge_repo.mark_not_current(v["doc_id"])

    # 激活目标版本
    await knowledge_repo.set_current(target_doc_id)
    return ApiResponse(data={"doc_id": target_doc_id, "version_no": version_no, "status": "restored"},
                       message=f"已恢复为 v{version_no}")